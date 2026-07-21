import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    doc.add_heading('Vultr Orchestration Migration Blueprint', 0)

    lines = content.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('> [!'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. '):
            doc.add_paragraph(line[3:], style='List Number')
        elif line.strip() == '':
            continue
        else:
            # Handle links: [text](url) -> we will just format them explicitly since python-docx doesn't do true hyperlinks easily without complex XML hacking.
            # We'll just do a regex replace to "text (url)"
            line = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', line)
            doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Saved to {docx_path}")

md_path = r"C:\Users\AviShemla\.gemini\antigravity\brain\cfc7c743-b169-4b8b-a5e8-6c10348c0c51\implementation_plan.md"
desktop_path = r"C:\Users\AviShemla\Desktop\Vultr_Migration_Blueprint.docx"

try:
    import docx
except ImportError:
    os.system("pip install python-docx")

markdown_to_docx(md_path, desktop_path)
